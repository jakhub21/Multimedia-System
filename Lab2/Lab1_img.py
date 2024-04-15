import numpy as np
import matplotlib.pyplot as plt
import cv2
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO

img1 = plt.imread('A3.png')
print(img1.dtype)
print(img1.shape)
print(np.min(img1), np.max(img1))


def imgToUInt8(img):
    if img.dtype == 'uint8':
        return img
    else:
        return (img*255).astype('uint8')


def imgToFloat(img):
    if img.dtype == 'float32':
        return img
    else:
        return img.astype('float32')/255


# plt.imshow(img1)
# plt.show()

R = img1[:, :, 0]
plt.imshow(R, cmap=plt.cm.gray)
G = img1[:, :, 1]
plt.imshow(G, cmap=plt.cm.gray)
B = img1[:, :, 2]
plt.imshow(B, cmap=plt.cm.gray)

Y1 = 0.299 * R + 0.587 * G + 0.114 * B
Y2 = 0.2126 * R + 0.7152 * G + 0.0722 * B

plt.imshow(Y1, cmap=plt.cm.gray)
# plt.show()
plt.imshow(Y2, cmap=plt.cm.gray)
# plt.show()

img_BGR = cv2.imread('A3.png')
img_RGB = cv2.cvtColor(img_BGR, cv2.COLOR_BGR2RGB)
img_BGR = cv2.cvtColor(img_RGB, cv2.COLOR_RGB2BGR)
plt.imshow(img_RGB)
# plt.show()
plt.imshow(img_BGR)
# plt.show()


# zadanie 2
# img = plt.imread('B01.png')
img = plt.imread('B02.jpg')
kopia1 = img.copy()


def show9img(kopia):
    R = kopia[:, :, 0]
    G = kopia[:, :, 1]
    B = kopia[:, :, 2]

    Y1 = 0.299 * R + 0.587 * G + 0.114 * B
    Y2 = 0.2126 * R + 0.7152 * G + 0.0722 * B

    imgR = np.copy(kopia)
    imgR[:, :, 1] = 0
    imgR[:, :, 2] = 0
    imgG = np.copy(kopia)
    imgG[:, :, 0] = 0
    imgG[:, :, 2] = 0
    imgB = np.copy(kopia)
    imgB[:, :, 0] = 0
    imgB[:, :, 1] = 0

    fig, axs = plt.subplots(3, 3, figsize=(10, 7))
    axs[0, 0].imshow(kopia)
    axs[0, 0].set_title('Originalny')
    axs[0, 1].imshow(Y1, cmap=plt.cm.gray)
    axs[0, 1].set_title('Y1')
    axs[0, 2].imshow(Y2, cmap=plt.cm.gray)
    axs[0, 2].set_title('Y2')
    axs[1, 0].imshow(R, cmap=plt.cm.gray)
    axs[1, 0].set_title('R')
    axs[1, 1].imshow(G, cmap=plt.cm.gray)
    axs[1, 1].set_title('G')
    axs[1, 2].imshow(B, cmap=plt.cm.gray)
    axs[1, 2].set_title('B')
    axs[2, 0].imshow(imgR)
    axs[2, 0].set_title('R')
    axs[2, 1].imshow(imgG)
    axs[2, 1].set_title('G')
    axs[2, 2].imshow(imgB)
    axs[2, 2].set_title('B')
    plt.show()

    memfile = BytesIO()  # tworzenie bufora
    fig.savefig(memfile)  # z zapis do bufora
    document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku
    memfile.close()

# show9img(kopia1)

# zadanie 3
fragment = img[0:200, 0:200]

df = pd.DataFrame()
df = pd.DataFrame(data={'Filename': ['B02.jpg'], 'Grayscale': [False],
                        'Fragments': [[[600, 700, 800, 900], [200, 300, 400, 500]]]})
print(df)

document = Document()
document.add_heading('Hubert Jakubiak LAB2', 0)

for index, row in df.iterrows():
    img = plt.imread(row['Filename'])
    if row['Grayscale']:
        # GS image - teraz nas nie intersuje
        pass
    else:
        # wyswietlenie subplota z funkcji
        document.add_heading(f"Oryginał", 2)
        show9img(img)

    if row['Fragments'] is not None:
        # mamy nie pustą listę fragmentów
        for f in row['Fragments']:
            document.add_heading(f"Fragment - height: {f[0]}, {f[2]} - width: {f[1]}, {f[3]}", 2)
            fragment = img[f[0]:f[2], f[1]:f[3]].copy()
            # tu wykonujesz operacje i inne wyświetlenia na fragmencie
            show9img(fragment)
document.save('report.docx')
