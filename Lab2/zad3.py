import numpy as np
import matplotlib.pyplot as plt
import cv2
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO

img = plt.imread('B02.jpg')
kopia1 = img.copy()


def show9img(kopia):
    R = kopia[:, :, 0]
    G = kopia[:, :, 1]
    B = kopia[:, :, 2]

    Y1 = 0.299 * R + 0.587 * G + 0.114 * B
    Y2 = 0.2126 * R + 0.7152 * G + 0.0722 * B

    imgR = np.copy(kopia)
    imgR[:, :, 1] = 0
    imgR[:, :, 2] = 0
    imgG = np.copy(kopia)
    imgG[:, :, 0] = 0
    imgG[:, :, 2] = 0
    imgB = np.copy(kopia)
    imgB[:, :, 0] = 0
    imgB[:, :, 1] = 0

    fig, axs = plt.subplots(3, 3, figsize=(10, 7))
    axs[0, 0].imshow(kopia)
    axs[0, 0].set_title('Originalny')
    axs[0, 1].imshow(Y1, cmap=plt.cm.gray)
    axs[0, 1].set_title('Y1')
    axs[0, 2].imshow(Y2, cmap=plt.cm.gray)
    axs[0, 2].set_title('Y2')
    axs[1, 0].imshow(R, cmap=plt.cm.gray)
    axs[1, 0].set_title('R')
    axs[1, 1].imshow(G, cmap=plt.cm.gray)
    axs[1, 1].set_title('G')
    axs[1, 2].imshow(B, cmap=plt.cm.gray)
    axs[1, 2].set_title('B')
    axs[2, 0].imshow(imgR)
    axs[2, 0].set_title('R')
    axs[2, 1].imshow(imgG)
    axs[2, 1].set_title('G')
    axs[2, 2].imshow(imgB)
    axs[2, 2].set_title('B')
    plt.show()

    memfile = BytesIO()  # tworzenie bufora
    fig.savefig(memfile)  # z zapis do bufora
    document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku
    memfile.close()

# show9img(kopia1)

# zadanie 3
fragment = img[0:200, 0:200]

df = pd.DataFrame()
df = pd.DataFrame(data={'Filename': ['B02.jpg'], 'Grayscale': [False],
                        'Fragments': [[[600, 700, 800, 900], [200, 300, 400, 500]]]})
print(df)

document = Document()
document.add_heading('Hubert Jakubiak LAB2', 0)

for index, row in df.iterrows():
    img = plt.imread(row['Filename'])
    if row['Grayscale']:
        # GS image - teraz nas nie intersuje
        pass
    else:
        # wyswietlenie subplota z funkcji
        show9img(img)

    if row['Fragments'] is not None:
        # mamy nie pustą listę fragmentów
        for f in row['Fragments']:
            fragment = img[f[0]:f[2], f[1]:f[3]].copy()
            # tu wykonujesz operacje i inne wyświetlenia na fragmencie
            show9img(fragment)


document.save('report.docx')