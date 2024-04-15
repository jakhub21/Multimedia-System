import numpy as np
import matplotlib.pyplot as plt
from scipy.interpolate import interp1d
import scipy
import soundfile as sf
import sounddevice as sd
import cv2
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO
from docx2pdf import convert


def kwant(data, bit):
    d = (2**bit)-1
    if np.issubdtype(data.dtype, np.floating):
        m = -1
        n = 1
    else:
        m = np.iinfo(data.dtype).min
        n = np.iinfo(data.dtype).max
    data_f = data.astype(float)
    data_f = (((np.round(((data_f - m)/(n-m))*d))/d)*(n-m))+m
    return data_f.astype(data.dtype)


def decimation(data, n, fs):
    new_data = data[::n]
    new_fs = fs//n
    return new_data, new_fs


def interpolation_lin(data, old_fs, new_fs):
    x = np.linspace(0, len(data)-1, len(data))
    f = interp1d(x, data, kind='linear')
    x_new = np.arange(0, len(data)-1, old_fs/new_fs)
    new_data = f(x_new)
    return new_data, new_fs


def interpolation_nonlin(data, old_fs, new_fs):
    x = np.linspace(0, len(data)-1, len(data))
    f = interp1d(x, data, kind='cubic')
    x_new = np.arange(0, len(data)-1, old_fs/new_fs)
    new_data = f(x_new)
    return new_data, new_fs


def plot_audio(signal, fs, time_margin=[0, 0.02], fsize=2**8, fig=None, axs=None):
    if fig is None:
        fig, axs = plt.subplots(2, 1, figsize=(10, 9))
    time = np.arange(0, signal.shape[0]) / fs
    axs[0].plot(time, signal)
    axs[0].set_xlim(time_margin)
    axs[0].set_title('Czas')
    axs[0].set_xlabel('Czas [s]')
    axs[0].set_ylabel('Amplituda')
    yf = scipy.fftpack.fft(signal, fsize)
    freq = np.arange(0, fs / 2, fs / fsize)
    axs[1].plot(freq, 20 * np.log10(np.abs(yf[:fsize // 2])) + np.finfo(np.float32).eps)
    axs[1].set_title('Częstotliwość')
    axs[1].set_xlabel('Częstotliwość [Hz]')
    axs[1].set_ylabel('Amplituda [dB]')
    max_idx = np.argmax(np.abs(yf[:fsize // 2]))
    peak_freq = np.arange(0, fs / 2, fs / fsize)[max_idx]
    peak_amp = 20 * np.log10(np.abs(yf[max_idx])) + np.finfo(np.float32).eps
    return fig, axs, peak_freq, peak_amp


# x1 = np.round(np.linspace(0, 255, 255, dtype=np.uint8))
# x2 = np.round(np.linspace(np.iinfo(np.int32).min, np.iinfo(np.int32).max, 1000, dtype=np.int32))
# x3 = np.linspace(-1, 1, 10000)
#
# plt.plot(x1, kwant(x1, 3))
# plt.show()

document = Document()
document.add_heading('Hubert Jakubiak LAB5', 0)

#zad1
files = ['SIN/sin_60Hz.wav', 'SIN/sin_440Hz.wav', 'SIN/sin_8000Hz.wav', 'SIN/sin_combined.wav']
bit = [4, 8, 16, 24]
n = [2, 4, 6, 10, 24]
fs_2 = [2000, 4000, 8000, 11999, 16000, 16953, 24000, 41000]
for file in files:
    for i, Bit in enumerate(bit):
        document.add_heading('Plik - {}'.format(file), 2)
        document.add_heading('Kwantyzacja, Bit {}'.format(Bit), 3)
        tmp = sf.read(file, dtype=np.int32)
        data = kwant(tmp[0], Bit)
        fs = tmp[1]
        fig, axs, peak_freq, peak_amp = plot_audio(data, fs)
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        document.add_paragraph('wartość losowa = {}'.format(np.random.rand(1)))
        document.add_paragraph('Peak freq = {}'.format(peak_freq))
        document.add_paragraph('Peak amp = {}'.format(peak_amp))

    for i, N in enumerate(n):
        document.add_heading('Plik - {}'.format(file), 2)
        document.add_heading('Decymacja, N {}'.format(N), 3)
        data, fs = sf.read(file, dtype=np.int32)
        data, fs = decimation(data, N, fs)
        fig, axs, peak_freq, peak_amp = plot_audio(data, fs)
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        document.add_paragraph('wartość losowa = {}'.format(np.random.rand(1)))
        document.add_paragraph('Peak freq = {}'.format(peak_freq))
        document.add_paragraph('Peak amp = {}'.format(peak_amp))

    for i, Fs in enumerate(fs_2):
        document.add_heading('Plik - {}'.format(file), 2)
        document.add_heading('Interpolacja liniowa, Fs {}'.format(Fs), 3)
        data, fs = sf.read(file, dtype=np.int32)
        data, fs = interpolation_lin(data, fs, Fs)
        fig, axs, peak_freq, peak_amp = plot_audio(data, fs)
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        document.add_paragraph('wartość losowa = {}'.format(np.random.rand(1)))
        document.add_paragraph('Peak freq = {}'.format(peak_freq))
        document.add_paragraph('Peak amp = {}'.format(peak_amp))

    for i, Fs in enumerate(fs_2):
        document.add_heading('Plik - {}'.format(file), 2)
        document.add_heading('Interpolacja nieliniowa, Fs {}'.format(Fs), 3)
        data, fs = sf.read(file, dtype=np.int32)
        data, fs = interpolation_nonlin(data, fs, Fs)
        fig, axs, peak_freq, peak_amp = plot_audio(data, fs)
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        document.add_paragraph('wartość losowa = {}'.format(np.random.rand(1)))
        document.add_paragraph('Peak freq = {}'.format(peak_freq))
        document.add_paragraph('Peak amp = {}'.format(peak_amp))


#zad2
Files = ['SING/sing_low1.wav', 'SING/sing_high1.wav', 'SING/sing_medium1.wav']
Bit_2 = [4, 8]
N_2 = [4, 6, 10, 24]
Fs_2 = [4000, 8000, 11999, 16000, 16953]
for file in Files:
    for i, Bit in enumerate(Bit_2):
        document.add_heading('Plik - {}'.format(file), 2)
        document.add_heading('Kwantyzacja, Bit {}'.format(Bit), 3)
        tmp = sf.read(file, dtype=np.int32)
        data = kwant(tmp[0], Bit)
        fs = tmp[1]
        sf.write(f'{file}Bit{Bit}.wav', data, fs)
        fig, axs, peak_freq, peak_amp = plot_audio(data, fs)
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        document.add_paragraph('wartość losowa = {}'.format(np.random.rand(1)))
        document.add_paragraph('Peak freq = {}'.format(peak_freq))
        document.add_paragraph('Peak amp = {}'.format(peak_amp))

    for i, N in enumerate(N_2):
        document.add_heading('Plik - {}'.format(file), 2)
        document.add_heading('Decymacja, N {}'.format(N), 3)
        data, fs = sf.read(file, dtype=np.int32)
        data, fs = decimation(data, N, fs)
        sf.write(f'{file}N{N}.wav', data, fs)
        fig, axs, peak_freq, peak_amp = plot_audio(data, fs)
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        document.add_paragraph('wartość losowa = {}'.format(np.random.rand(1)))
        document.add_paragraph('Peak freq = {}'.format(peak_freq))
        document.add_paragraph('Peak amp = {}'.format(peak_amp))

    for i, Fs in enumerate(Fs_2):
        document.add_heading('Plik - {}'.format(file), 2)
        document.add_heading('Interpolacja liniowa, Fs {}'.format(Fs), 3)
        data, fs = sf.read(file, dtype=np.int32)
        data, fs = interpolation_lin(data, fs, Fs)
        sf.write(f'{file}Fs{Fs}L.wav', data, fs)
        fig, axs, peak_freq, peak_amp = plot_audio(data, fs)
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        document.add_paragraph('wartość losowa = {}'.format(np.random.rand(1)))
        document.add_paragraph('Peak freq = {}'.format(peak_freq))
        document.add_paragraph('Peak amp = {}'.format(peak_amp))

    for i, Fs in enumerate(Fs_2):
        document.add_heading('Plik - {}'.format(file), 2)
        document.add_heading('Interpolacja nieliniowa, Fs {}'.format(Fs), 3)
        data, fs = sf.read(file, dtype=np.int32)
        data, fs = interpolation_nonlin(data, fs, Fs)
        sf.write(f'{file}Fs{Fs}NL.wav', data, fs)
        fig, axs, peak_freq, peak_amp = plot_audio(data, fs)
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        document.add_paragraph('wartość losowa = {}'.format(np.random.rand(1)))
        document.add_paragraph('Peak freq = {}'.format(peak_freq))
        document.add_paragraph('Peak amp = {}'.format(peak_amp))

document.save('report.docx')

