#funkcja ma przyjmowac argumenty (obraz, skala)
import numpy as np
import matplotlib.pyplot as plt
import cv2
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO
from docx2pdf import convert

# img = np.zeros((3, 3, 3), dtype=np.uint8)
# img[1, 1, :] = 255


def metoda_najblizszego_sasiada(img, scale):
    height, width = img.shape[:2]
    new_height = np.ceil(height * scale).astype(int)
    new_width = np.ceil(width * scale).astype(int)
    new_img = np.zeros((new_height, new_width, 3), dtype=np.uint8)
    x = np.linspace(0, height - 1, new_height)
    y = np.linspace(0, width - 1, new_width)
    for i in range(0, new_height):
        for j in range(0, new_width):
            new_img[i, j] = img[np.round(x[i]).astype(int), np.round(y[j]).astype(int)]
    edges = cv2.Canny(new_img, 100, 200)  # Dodanie wykrywania krawędzi
    return new_img, edges


def interpolacja_dwuliniowa(img, scale):
    height, width = img.shape[:2]
    new_height = np.ceil(height * scale).astype(int)
    new_width = np.ceil(width * scale).astype(int)
    new_img = np.zeros((new_height, new_width, 3), dtype=np.uint8)
    x = np.linspace(0, height - 1, new_height)
    y = np.linspace(0, width - 1, new_width)
    for i in range(0, new_height):
        for j in range(0, new_width):
            xl = np.floor(x[i]).astype(int)
            xr = np.ceil(x[i]).astype(int)
            yl = np.floor(y[j]).astype(int)
            yr = np.ceil(y[j]).astype(int)
            dy = y[j] - yl
            dx = x[i] - xl
            new_img[i, j] = (1 - dx) * (1 - dy) * img[xl, yl] + (1 - dx) * dy * img[xl, yr] + dx * (1 - dy) * img[
                xr, yl] + dx * dy * img[xr, yr]
    edges = cv2.Canny(new_img, 100, 200)  # Dodanie wykrywania krawędzi
    return new_img, edges


def zmniejszanie_sredia(img, scale):
    height, width = img.shape[:2]
    new_height = int(np.ceil(height * scale))
    new_width = int(np.ceil(width * scale))
    new_img = np.zeros((new_height, new_width, 3), dtype=np.uint8)
    x = np.linspace(0, width - 1, new_width)
    y = np.linspace(0, height - 1, new_height)
    for i in range(new_height):
        for j in range(new_width):
            xl = np.round(x[j] + np.arange(-3, 4)).astype(int)
            yl = np.round(y[i] + np.arange(-3, 4)).astype(int)
            xl = (xl[(xl >= 0) & (xl < width)])
            yl = (yl[(yl >= 0) & (yl < height)])
            xx, yy = np.meshgrid(xl, yl)
            if img.ndim > 2:
                for z in range(3):
                    new_img[i, j, z] = np.mean(img[yy, xx, z])
            else:
                new_img[i, j] = np.mean(img[yy, xx])
    edges = cv2.Canny(new_img, 100, 200)  # Dodanie wykrywania krawędzi
    return new_img, edges


def zmniejszanie_mediana(img, scale):
    height, width = img.shape[:2]
    new_height = int(np.ceil(height * scale))
    new_width = int(np.ceil(width * scale))
    new_img = np.zeros((new_height, new_width, 3), dtype=np.uint8)
    x = np.linspace(0, width - 1, new_width)
    y = np.linspace(0, height - 1, new_height)
    for i in range(new_height):
        for j in range(new_width):
            xl = np.round(x[j] + np.arange(-3, 4)).astype(int)
            yl = np.round(y[i] + np.arange(-3, 4)).astype(int)
            xl = (xl[(xl >= 0) & (xl < width)])
            yl = (yl[(yl >= 0) & (yl < height)])
            xx, yy = np.meshgrid(xl, yl)
            if img.ndim > 2:
                for z in range(3):
                    new_img[i, j, z] = np.median(img[yy, xx, z], axis=(0, 1))
            else:
                new_img[i, j] = np.median(img[yy, xx], axis=(0, 1))
    edges = cv2.Canny(new_img, 100, 200)  # Dodanie wykrywania krawędzi
    return new_img, edges


def zmniejszanie_wazona(img, scale):
    height, width = img.shape[:2]
    new_height = int(np.ceil(height * scale))
    new_width = int(np.ceil(width * scale))
    new_img = np.zeros((new_height, new_width, 3), dtype=np.uint8)
    x = np.linspace(0, width - 1, new_width)
    y = np.linspace(0, height - 1, new_height)
    for i in range(new_height):
        for j in range(new_width):
            xl = np.round(x[j] + np.arange(-3, 4)).astype(int)
            yl = np.round(y[i] + np.arange(-3, 4)).astype(int)
            xl = (xl[(xl >= 0) & (xl < width)])
            yl = (yl[(yl >= 0) & (yl < height)])
            xx, yy = np.meshgrid(xl, yl)
            wagi = np.random.rand(len(yl), len(xl))
            if img.ndim > 2:
                for z in range(3):
                    new_img[i, j, z] = np.sum(img[yy, xx, z] * wagi)/np.sum(wagi)
            else:
                new_img[i, j] = np.sum(img[yy, xx] * wagi)/np.sum(wagi)
    edges = cv2.Canny(new_img, 100, 200)  # Dodanie wykrywania krawędzi
    return new_img, edges





df = pd.DataFrame()
df = pd.DataFrame(data={'FilenameB': ['BIG_0001.jpg', 'BIG_0002.jpg', 'BIG_0004.png'],
                        'Fragments': [((300, 600), (300, 600)), ((400, 800), (400, 800)), ((400, 600), (400, 600))],
                        'Scale': [0.1, 0.2, 0.3]})

df2 = pd.DataFrame()
df2 = pd.DataFrame(data={'FilenameS': ['SMALL_0002.png', 'SMALL_0005.jpg', 'SMALL_0008.jpg'],
                         'Fragments2': [((0, 50), (0, 50)), ((0, 100), (0, 100)), ((20, 80), (20, 80))],
                         'Scale': [3, 4, 5]})

document = Document()
document.add_heading('Hubert Jakubiak LAB3', 0)

document.add_heading('Pomniejszanie zdjecia', 1)
for index, row in df.iterrows():
    for img in (df["FilenameB"]):
        img = cv2.imread(img)
        img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
        for scale in (df["Scale"]):
            document.add_heading('Skala - {}'.format(scale), 2)
            if row['Fragments'] is not None:
                fragment = img[row['Fragments'][0][0]:row['Fragments'][0][1], row['Fragments'][1][0]:row['Fragments'][1][1]]
                document.add_heading('Fragment - {}'.format(row['Fragments']), 3)
                fig, axs = plt.subplots(6, 2, figsize=(7, 9))

                axs[0][0].imshow(fragment, cmap='gray')
                axs[0][0].set_title('Originalny')
                axs[0][0].axis('off')
                edges = cv2.Canny(fragment, 100, 200)
                axs[0][1].imshow(edges, cmap='gray')
                axs[0][1].set_title('Originalny')
                axs[0][1].axis('off')
                new_fragment, edges = metoda_najblizszego_sasiada(fragment, scale)
                axs[1][0].imshow(new_fragment, cmap='gray')
                axs[1][0].set_title('Najbliższy sąsiad')
                axs[1][0].axis('off')
                axs[1][1].imshow(edges, cmap='gray')
                axs[1][1].set_title('Najbliższy sąsiad')
                axs[1][1].axis('off')
                new_fragment, edges = interpolacja_dwuliniowa(fragment, scale)
                axs[2][0].imshow(new_fragment, cmap='gray')
                axs[2][0].set_title('Interpolacja dwuliniowa')
                axs[2][0].axis('off')
                axs[2][1].imshow(edges, cmap='gray')
                axs[2][1].set_title('Interpolacja dwuliniowa')
                axs[2][1].axis('off')
                new_fragment, edges = zmniejszanie_sredia(fragment, scale)
                axs[3][0].imshow(new_fragment, cmap='gray')
                axs[3][0].set_title('Średnia')
                axs[3][0].axis('off')
                axs[3][1].imshow(edges, cmap='gray')
                axs[3][1].set_title('Średnia')
                axs[3][1].axis('off')
                new_fragment, edges = zmniejszanie_mediana(fragment, scale)
                axs[4][0].imshow(new_fragment, cmap='gray')
                axs[4][0].set_title('Mediana')
                axs[4][0].axis('off')
                axs[4][1].imshow(edges, cmap='gray')
                axs[4][1].set_title('Mediana')
                axs[4][1].axis('off')
                new_fragment, edges = zmniejszanie_wazona(fragment, scale)
                axs[5][0].imshow(new_fragment, cmap='gray')
                axs[5][0].set_title('Ważona')
                axs[5][0].axis('off')
                axs[5][1].imshow(edges, cmap='gray')
                axs[5][1].set_title('Ważona')
                axs[5][1].axis('off')
                # plt.show()
                memfile = BytesIO()
                fig.savefig(memfile)
                document.add_picture(memfile, width=Inches(6))
                memfile.close()


document.add_heading('Powiekszanie zdjecia', 1)
for index, row2 in df2.iterrows():
    for img2 in (df2["FilenameS"]):
        img2 = cv2.imread(img2)
        img2 = cv2.cvtColor(img2, cv2.COLOR_BGR2RGB)
        for scale in (df2["Scale"]):
            document.add_heading('Skala - {}'.format(scale), 2)
            if row2['Fragments2'] is not None:
                fragment = img2[row2['Fragments2'][0][0]:row2['Fragments2'][0][1], row2['Fragments2'][1][0]:row2['Fragments2'][1][1]]
                document.add_heading('Fragment - {}'.format(row2['Fragments2']), 3)
                fig, axs = plt.subplots(2, 3, figsize=(10, 6))
                axs[0][0].imshow(fragment, cmap='gray')
                axs[0][0].set_title('Originalny')
                edges = cv2.Canny(fragment, 100, 200)
                axs[1][0].imshow(edges, cmap='gray')
                new_fragment, edges = metoda_najblizszego_sasiada(fragment, 3)
                axs[0][1].imshow(new_fragment, cmap='gray')
                axs[0][1].set_title('Najbliższy sąsiad')
                axs[1][1].imshow(edges, cmap='gray')
                axs[1][1].set_title('Najbliższy sąsiad')
                new_fragment, edges = interpolacja_dwuliniowa(fragment, 3)
                axs[0][2].imshow(new_fragment, cmap='gray')
                axs[0][2].set_title('Interpolacja dwuliniowa')
                axs[1][2].imshow(edges, cmap='gray')
                axs[1][2].set_title('Interpolacja dwuliniowa')
                # plt.show()
                memfile = BytesIO()
                fig.savefig(memfile)
                document.add_picture(memfile, width=Inches(6))
                memfile.close()
document.save('report.docx')
# # # convert("report.docx")
# #
# #
# #
