import numpy as np
import matplotlib.pyplot as plt
from scipy.interpolate import interp1d
import scipy
import soundfile as sf
import sounddevice as sd
import cv2
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO
from docx2pdf import convert
import sys


def get_size(obj, seen=None):
    """Recursively finds size of objects"""
    size = sys.getsizeof(obj)
    if seen is None:
        seen = set()
    obj_id = id(obj)
    if obj_id in seen:
        return 0
    # Important mark as seen *before* entering recursion to gracefully handle
    # self-referential objects
    seen.add(obj_id)
    if isinstance(obj, np.ndarray):
        size = obj.nbytes
    elif isinstance(obj, dict):
        size += sum([get_size(v, seen) for v in obj.values()])
        size += sum([get_size(k, seen) for k in obj.keys()])
    elif hasattr(obj, '__dict__'):
        size += get_size(obj.__dict__, seen)
    elif hasattr(obj, '__iter__') and not isinstance(obj, (str, bytes, bytearray)):
        size += sum([get_size(i, seen) for i in obj])
    return size


def help_fun_same(text):
    counter = 1
    for i in range(len(text) - 1):
        counter += 1
        if text[i] != text[i + 1]:
            counter -= 1
            break
    return counter


def help_fun_different(text):
    counter = 1
    for i in range(len(text) - 1):
        if text[i] == text[i + 1]:
            counter -= 1
            break
        counter += 1
    return counter


def rle_encode(text):
    x = np.array([len(text.shape)])
    x = np.concatenate([x, text.shape])
    text = text.flatten()
    encoded = np.zeros(np.prod(text.shape)*2)
    encoded = encoded.astype(int)
    i = 0
    j = 0
    while i < len(text):
        counter = help_fun_same(text[i:])
        encoded[j] = counter
        encoded[j+1] = text[i]
        i += counter
        j += 2
    encoded = encoded[:j]
    encoded = np.concatenate([x, encoded])
    return encoded


def rle_decode(text):
    shape = text[1:int(text[0] + 1)]
    text = text[int(text[0] + 1):]
    decoded = np.zeros(np.prod(shape))
    decoded = decoded.astype(int)
    i = 0
    j = 0
    while i < len(text):
        for k in range(text[i]):
            decoded[j] = text[i+1]
            j += 1
        i += 2
    decoded = np.reshape(decoded, shape)
    return decoded


def byte_run_encode(text):
    x = np.array([len(text.shape)])
    x = np.concatenate([x, text.shape])
    text = text.flatten()
    encoded = np.zeros(np.prod(text.shape) * 2)
    encoded = encoded.astype(int)
    i = 0
    j = 0
    while i < len(text):
        counter = help_fun_same(text[i:])
        if counter > 1:
            encoded[j] = -counter + 1
            encoded[j + 1] = text[i]
            i += counter
            j += 2
        else:
            counter = help_fun_different(text[i:])
            encoded[j] = counter - 1
            for k in range(counter):
                encoded[j + k + 1] = text[i + k]
            i += counter
            j += counter + 1
    encoded = encoded[:j]
    encoded = np.concatenate([x, encoded])
    return encoded


def byte_run_decode(text):
    shape = text[1:int(text[0] + 1)]
    text = text[int(text[0] + 1):]
    decoded = np.zeros(np.prod(shape))
    decoded = decoded.astype(int)
    i = 0
    j = 0
    while i < len(text):
        if text[i] < 0:
            for k in range(-text[i] + 1):
                decoded[j] = text[i + 1]
                j += 1
            i += 2
        else:
            for k in range(text[i] + 1):
                decoded[j] = text[i + k + 1]
                j += 1
            i += text[i] + 2
    decoded = np.reshape(decoded, shape)
    return decoded


t1 = np.array([1,1,1,1,2,1,1,1,1,2,1,1,1,1]).astype(int)
t2 = np.array([1,2,3,1,2,3,1,2,3])
t3 = np.array([5,1,5,1,5,5,1,1,5,5,1,1,5])
t4 = np.array([-1,-1,-1,-5,-5,-3,-4,-2,1,2,2,1])
t5 = np.zeros((1,520))
t6 = np.arange(0,521,1)
t7 = np.eye(7)
t8 = np.dstack([np.eye(7),np.eye(7),np.eye(7)])
t9 = np.ones((1,1,1,1,1,1,10))

t = [t1, t2, t3, t4, t5, t6, t7, t8, t9]
#RLE
# for i in t:
#     encode = rle_encode(i)
#     print(encode)
#     decode = rle_decode(encode)
#     print(decode)
#     print(i)
#     print("\n")
#
#Byte Run
# for i in t:
#     encode = byte_run_encode(i)
#     print(encode)
#     decode = byte_run_decode(encode)
#     print(decode)
#     print(i)
#     print("\n")

document = Document()
document.add_heading('Hubert Jakubiak LAB6', 0)

df = pd.DataFrame()
df = pd.DataFrame(data={'Filename': ['rysunek_techniczny.png', 'formularz.png', 'kolor.jpg']})

for img in (df["Filename"]):
    document.add_heading('IMG = {}'.format(img), level=1)
    img = cv2.imread(img)
    img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    fig, axs = plt.subplots(3, 1, figsize=(7, 7))
    axs[0].imshow(img, cmap='gray')
    axs[0].set_title('Original Image')
    axs[0].axis('off')
    encode_rle = rle_encode(img)
    size_rle = get_size(encode_rle)
    decode_rle = rle_decode(encode_rle)

    axs[1].imshow(decode_rle, cmap='gray')
    axs[1].set_title('RLE')
    axs[1].axis('off')
    encode_byte_run = byte_run_encode(img)
    size_byte_run = get_size(encode_byte_run)
    decode_byte_run = byte_run_decode(encode_byte_run)

    axs[2].imshow(decode_byte_run, cmap='gray')
    axs[2].set_title('Byte Run')
    axs[2].axis('off')
    # plt.show()
    cr_rle = abs(get_size(img))/abs(size_rle)
    cr_byte = abs(get_size(img))/abs(size_byte_run)
    pr_rle = abs(size_rle)/abs(get_size(img))*100
    pr_byte = abs(size_byte_run)/abs(get_size(img))*100
    document.add_paragraph('RLE Compression Ratio: {}'.format(cr_rle))
    document.add_paragraph('Byte Run Compression Ratio: {}'.format(cr_byte))
    document.add_paragraph('RLE Compression Percentage: {}'.format(pr_rle))
    document.add_paragraph('Byte Run Compression Percentage: {}'.format(pr_byte))
    memfile = BytesIO()
    fig.savefig(memfile)
    document.add_picture(memfile, width=Inches(6))
    memfile.close()
document.save('report.docx')




    

