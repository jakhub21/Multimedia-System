import numpy as np
import matplotlib.pyplot as plt
import cv2
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO
from docx2pdf import convert


def imgToFloat(img):
    if img.dtype == 'float32':
        return img
    else:
        return img.astype('float32')/255


def color_fit(pixel, pallet):
    if len(pallet.shape) > 1:
        return pallet[np.argmin(np.linalg.norm(pallet - pixel, axis=1))]
    else:
        return pallet


def kwant_colorFit(img, Pallet):
    img = imgToFloat(img)
    out_img = img.copy()
    for w in range(img.shape[0]):
        for k in range(img.shape[1]):
            out_img[w, k] = color_fit(img[w, k], Pallet)
    return out_img


def dithering_random(img):
    img = imgToFloat(img)
    if len(img.shape) > 2:
        img = img[:, :, 0].copy()
    else:
        img = img.copy()
    r = np.random.rand(img.shape[0], img.shape[1])
    img = img >= r
    return img * 1.0


def dithering_ordered(img, pallet):
    img = imgToFloat(img)
    M = np.array([
        [0, 8, 2, 10],
        [12, 4, 14, 6],
        [3, 11, 1, 9],
        [15, 7, 13, 5]
    ])
    n = 2

    if len(img.shape) > 2:
        out_img = img[:, :, :3].copy()
    else:
        out_img = img.copy()
    Mpre = (M + 1) / (2 * n) ** 2 - 0.5
    for i in range(img.shape[0]):
        for j in range(img.shape[1]):
            out_img[i][j] = color_fit(img[i][j] + Mpre[i % (2*n), j % (2*n)], pallet)
    return out_img


def dithering_floyd_steinberg(img, pallet):
    img = imgToFloat(img)
    if len(img.shape) > 2:
        img = img[:, :, :3].copy()
    else:
        img = img.copy()
    for i in range(img.shape[0]):
        for j in range(img.shape[1]):
            old_pixel = img[i, j].copy()
            new_pixel = color_fit(old_pixel, pallet)
            img[i, j] = new_pixel
            error = old_pixel - new_pixel
            if j < img.shape[1] - 1:
                img[i, j + 1] += (error * 7 / 16)
            if i < img.shape[0] - 1 and j > 0:
                img[i + 1, j - 1] += (error * 3 / 16)
            if i < img.shape[0] - 1:
                img[i + 1, j] += (error * 5 / 16)
            if i < img.shape[0] - 1 and j < img.shape[1] - 1:
                img[i + 1, j + 1] += (error * 1 / 16)
    return img


pallet1 = np.expand_dims(np.linspace(0.0, 1.0, 2), axis=1)
pallet2 = np.expand_dims(np.linspace(0.0, 1.0, 4), axis=1)
pallet4 = np.expand_dims(np.linspace(0.0, 1.0, 16), axis=1)

pallet8 = np.array([
        [0.0, 0.0, 0.0,],
        [0.0, 0.0, 1.0,],
        [0.0, 1.0, 0.0,],
        [0.0, 1.0, 1.0,],
        [1.0, 0.0, 0.0,],
        [1.0, 0.0, 1.0,],
        [1.0, 1.0, 0.0,],
        [1.0, 1.0, 1.0,],
])

pallet16 = np.array([
        [0.0, 0.0, 0.0,],
        [0.0, 1.0, 1.0,],
        [0.0, 0.0, 1.0,],
        [1.0, 0.0, 1.0,],
        [0.0, 0.5, 0.0,],
        [0.5, 0.5, 0.5,],
        [0.0, 1.0, 0.0,],
        [0.5, 0.0, 0.0,],
        [0.0, 0.0, 0.5,],
        [0.5, 0.5, 0.0,],
        [0.5, 0.0, 0.5,],
        [1.0, 0.0, 0.0,],
        [0.75, 0.75, 0.75,],
        [0.0, 0.5, 0.5,],
        [1.0, 1.0, 1.0,],
        [1.0, 1.0, 0.0,]
])

# fig, axs = plt.subplots(2, 3, figsize=(10, 7))
# img_float = cv2.imread('IMG_SMALL/SMALL_0006.jpg')
# img_float = cv2.cvtColor(img_float, cv2.COLOR_BGR2RGB)
# axs[0][0].imshow(img_float, cmap='gray')
# axs[0][0].set_title('Originalny')
# axs[0][0].axis('off')
# axs[0][1].imshow(kwant_colorFit(img_float, pallet16), cmap='gray')
# axs[0][1].set_title('Kwantyzacja 2')
# axs[0][1].axis('off')
# axs[0][2].imshow(dithering_random(img_float), cmap='gray')
# axs[0][2].set_title('Dithering losowy 2')
# axs[0][2].axis('off')
# axs[1][0].imshow(dithering_ordered(img_float, pallet16), cmap='gray')
# axs[1][0].set_title('Dithering uporządkowany 2')
# axs[1][0].axis('off')
# axs[1][1].imshow(dithering_floyd_steinberg(img_float, pallet16), cmap='gray')
# axs[1][1].set_title('Dithering Floyd-Steinberg 2')
# axs[1][1].axis('off')
# plt.show()

document = Document()
document.add_heading('Hubert Jakubiak LAB4', 0)

df = pd.DataFrame()
df = pd.DataFrame(data={'FilenameGS': ['IMG_GS/GS_0001.tif', 'IMG_GS/GS_0002.png', 'IMG_GS/GS_0003.png']})
pallets = [pallet1, pallet2, pallet4]

df2 = pd.DataFrame(data={'FilenameSMALL': ['IMG_SMALL/SMALL_0004.jpg', 'IMG_SMALL/SMALL_0006.jpg', 'IMG_SMALL/SMALL_0007.jpg', 'IMG_SMALL/SMALL_0009.jpg']})
pallets2 = [pallet8, pallet16]

for img in (df["FilenameGS"]):
    document.add_heading('GS', level=1)
    document.add_heading('IMG = {}'.format(img), level=2)
    img = cv2.imread(img)
    img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    for pallet in pallets:
        document.add_heading('Pallet = {}'.format(len(pallet)), level=3)
        fig, axs = plt.subplots(2, 3, figsize=(10, 7))
        axs[0][0].imshow(img, cmap='gray')
        axs[0][0].set_title('Originalny')
        axs[0][0].axis('off')
        axs[0][1].imshow(kwant_colorFit(img, pallet), cmap='gray')
        axs[0][1].set_title('Kwantyzacja')
        axs[0][1].axis('off')
        axs[0][2].imshow(dithering_random(img), cmap='gray')
        axs[0][2].set_title('Dithering losowy')
        axs[0][2].axis('off')
        axs[1][0].imshow(dithering_ordered(img, pallet), cmap='gray')
        axs[1][0].set_title('Dithering uporządkowany')
        axs[1][0].axis('off')
        axs[1][1].imshow(dithering_floyd_steinberg(img, pallet), cmap='gray')
        axs[1][1].set_title('Dithering Floyd-Steinberg')
        axs[1][1].axis('off')
        # plt.show()
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()

for img in (df2["FilenameSMALL"]):
    document.add_heading('SMALL', level=1)
    document.add_heading('IMG = {}'.format(img), level=2)
    img = cv2.imread(img)
    img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    for pallet in pallets2:
        document.add_heading('Pallet = {}'.format(len(pallet)), level=3)
        fig, axs = plt.subplots(2, 2, figsize=(10, 7))
        axs[0][0].imshow(img, cmap='gray')
        axs[0][0].set_title('Originalny')
        axs[0][0].axis('off')
        axs[0][1].imshow(kwant_colorFit(img, pallet), cmap='gray')
        axs[0][1].set_title('Kwantyzacja')
        axs[0][1].axis('off')
        axs[1][0].imshow(dithering_ordered(img, pallet), cmap='gray')
        axs[1][0].set_title('Dithering uporządkowany')
        axs[1][0].axis('off')
        axs[1][1].imshow(dithering_floyd_steinberg(img, pallet), cmap='gray')
        axs[1][1].set_title('Dithering Floyd-Steinberg')
        axs[1][1].axis('off')
        # plt.show()
        memfile = BytesIO()
        fig.savefig(memfile)
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
document.save('report.docx')

# img = plt.imread('IMG_GS/GS_0002.png')
# print(np.unique(dithering_random(img.copy(), np.linspace(0, 1, 2).reshape(2, 1))).size)


