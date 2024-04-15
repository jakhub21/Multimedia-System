from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
import soundfile as sf
import scipy.fftpack
from io import BytesIO


data, fs = sf.read('sin_440Hz.wav', dtype=np.int32)


def plot_audio(signal, fs, time_margin=[0, 0.02], fsize=2**8, fig=None, axs=None):
    if fig is None:
        fig, axs = plt.subplots(2, 1, figsize=(10, 7))
    time = np.arange(0, signal.shape[0]) / fs
    axs[0].plot(time, signal)
    axs[0].set_xlim(time_margin)
    axs[0].set_title('Czas')
    axs[0].set_xlabel('Czas [s]')
    axs[0].set_ylabel('Amplituda')
    yf = scipy.fftpack.fft(signal, fsize)
    freq = np.arange(0, fs / 2, fs / fsize)
    axs[1].plot(freq, 20 * np.log10(np.abs(yf[:fsize // 2])+ 1e-10))
    axs[1].set_title('Częstotliwość')
    axs[1].set_xlabel('Częstotliwość [Hz]')
    axs[1].set_ylabel('Amplituda [dB]')
    max_idx = np.argmax(np.abs(yf[:fsize // 2]))
    peak_freq = np.arange(0, fs / 2, fs / fsize)[max_idx]
    peak_amp = 20 * np.log10(np.abs(yf[max_idx]) + 1e-10)
    return fig, axs, peak_freq, peak_amp

fig, axs, peak_freq, peak_amp = plot_audio(data, fs)
plt.show()



document = Document()
document.add_heading('Hubert Jakubiak LAB1', 0)  # tworzenie nagłówków druga wartość to poziom nagłówka

files = ['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav']
fsize = [2**8, 2**12, 2**16]
Margins = [[0, 0.02], [0.133, 0.155]]
for file in files:
    document.add_heading('Plik - {}'.format(file), 2)
    for i, Margin in enumerate(Margins):
        document.add_heading('Time margin {}'.format(Margin), 3)  # nagłówek sekcji, mozę być poziom wyżej
        fig, axs = plt.subplots(2, 1, figsize=(10, 7))  # tworzenie plota

        ############################################################
        # Tu wykonujesz jakieś funkcje i rysujesz wykresy
        data, fs = sf.read(file, dtype=np.int32)
        fig, axs, peak_freq, peak_amp = plot_audio(data, fs, time_margin=Margin, fsize=fsize, fig=fig, axs=axs)
        ############################################################

        fig.suptitle('Time margin {}'.format(Margin))  # Tytuł wykresu
        fig.tight_layout(pad=1.5)  # poprawa czytelności
        memfile = BytesIO()  # tworzenie bufora
        fig.savefig(memfile)  # z zapis do bufora

        document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku

        memfile.close()
        ############################################################
        # Tu dodajesz dane tekstowe - wartosci, wyjscie funkcji ect.
        document.add_paragraph('wartość losowa = {}'.format(np.random.rand(1)))
        document.add_paragraph('Peak freq = {}'.format(peak_freq))
        document.add_paragraph('Peak amp = {}'.format(peak_amp))

        ############################################################

document.save('report.docx')  # zapis do pliku