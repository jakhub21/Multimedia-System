from docx import Document
from docx.shared import Inches
import sounddevice as sd
import soundfile as sf
import matplotlib.pyplot as plt
import numpy as np
import scipy.fftpack
from io import BytesIO

data, fs = sf.read('sound1.wav', dtype='float32')

print(data.dtype)
print(data.shape)
# print(data)

sd.play(data, fs)
status = sd.wait()

sf.write('sound1.wav', data, fs)

plt.subplot(3, 1, 1)
plt.plot(data[:, 0])
plt.subplot(3, 1, 2)
plt.plot(data[:, 1])
plt.subplot(3, 1, 3)
plt.plot(np.mean(data, axis=1))
plt.show()

data, fs = sf.read('sin_440Hz.wav', dtype=np.int32)
plt.figure()
plt.subplot(2,1,1)
plt.plot(np.arange(0,data.shape[0])/fs,data)

plt.subplot(2,1,2)
yf = scipy.fftpack.fft(data)
plt.plot(np.arange(0,fs,1.0*fs/(yf.size)),np.abs(yf))
plt.show()

fsize=2**8
plt.figure()
plt.subplot(2,1,1)
plt.plot(np.arange(0,data.shape[0])/fs,data)
plt.subplot(2,1,2)
yf = scipy.fftpack.fft(data,fsize)
plt.plot(np.arange(0,fs/2,fs/fsize),20*np.log10( np.abs(yf[:fsize//2])))
plt.show()

data, fs = sf.read('sin_440Hz.wav', dtype=np.int32)


def plot_audio(signal, fs, time_margin=[0, 0.02], fsize=2**8, fig=None, axs=None):
    if fig is None:
        fig, axs = plt.subplots(2, 1, figsize=(10, 7))
    time = np.arange(0, signal.shape[0]) / fs
    axs[0].plot(time, signal)
    axs[0].set_xlim(time_margin)
    axs[0].set_title('Czas')
    axs[0].set_xlabel('Czas [s]')
    axs[0].set_ylabel('Amplituda')
    yf = scipy.fftpack.fft(signal, fsize)
    freq = np.arange(0, fs / 2, fs / fsize)
    axs[1].plot(freq, 20 * np.log10(np.abs(yf[:fsize // 2])))
    axs[1].set_title('Częstotliwość')
    axs[1].set_xlabel('Częstotliwość [Hz]')
    axs[1].set_ylabel('Amplituda [dB]')
    max_idx = np.argmax(np.abs(yf[:fsize // 2]))
    peak_freq = np.arange(0, fs / 2, fs / fsize)[max_idx]
    peak_amp = 20 * np.log10(np.abs(yf[max_idx]))
    return fig, axs, peak_freq, peak_amp

fig, axs, peak_freq, peak_amp = plot_audio(data, fs)
plt.show()



document = Document()
document.add_heading('Hubert Jakubiak LAB1', 0)  # tworzenie nagłówków druga wartość to poziom nagłówka

files = ['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav']
fsize = [2**8, 2**12, 2**16]
for file in files:
    document.add_heading('Plik - {}'.format(file), 2)
    for i, Fsize in enumerate(fsize):
        document.add_heading('Fsize {}'.format(Fsize), 3)  # nagłówek sekcji, mozę być poziom wyżej
        fig, axs = plt.subplots(2, 1, figsize=(10, 7))  # tworzenie plota

        ############################################################
        # Tu wykonujesz jakieś funkcje i rysujesz wykresy
        data, fs = sf.read(file, dtype=np.int32)
        fig, axs, peak_freq, peak_amp = plot_audio(data, fs, fsize=Fsize, fig=fig, axs=axs)
        ############################################################

        fig.suptitle('Fsize {}'.format(Fsize))  # Tytuł wykresu
        fig.tight_layout(pad=1.5)  # poprawa czytelności
        memfile = BytesIO()  # tworzenie bufora
        fig.savefig(memfile)  # z zapis do bufora

        document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku

        memfile.close()
        ############################################################
        # Tu dodajesz dane tekstowe - wartosci, wyjscie funkcji ect.
        document.add_paragraph('wartość losowa = {}'.format(np.random.rand(1)))
        document.add_paragraph('Peak freq = {}'.format(peak_freq))
        document.add_paragraph('Peak amp = {}'.format(peak_amp))

        ############################################################

document.save('report.docx')  # zapis do pliku